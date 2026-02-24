import os
import re
import shutil
from docx import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LC_Document

# 1. Загрузка текста из Word
def load_docx(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл {file_path} не найден!")
    
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        # Сохраняем текст, убирая лишние пробелы по краям
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Объединяем через перевод строки для корректной работы Regex
    return "\n".join(full_text)

# 2. Структурная нарезка по статьям
def split_by_articles(text):
    # ПРАВИЛО: Статья начинается с новой строки, слова "Статья", номера и ТОЧКИ.
    # Используем positive lookahead (?=...), чтобы не "съедать" заголовок при разделении.
    # Флаг re.MULTILINE заставляет ^ искать начало каждой строки.
    article_pattern = r'(?=\nСтатья\s+\d+\.)'
    
    # Добавляем \n в начало текста, чтобы первая статья тоже попала под паттерн
    segments = re.split(article_pattern, "\n" + text, flags=re.MULTILINE)
    
    docs = []
    for segment in segments:
        content = segment.strip()
        if not content:
            continue
            
        # Извлекаем номер статьи для метаданных (например, "Статья 560.")
        # Это поможет модели ориентироваться в базе
        match = re.search(r'Статья\s+(\d+)\.', content)
        metadata = {"source": "TK_Uz_Clean.docx"}
        
        if match:
            metadata["article_number"] = int(match.group(1))
            metadata["type"] = "article"
        else:
            metadata["type"] = "intro_or_other"

        docs.append(LC_Document(page_content=content, metadata=metadata))
    
    return docs

def rebuild_database():
    # Указываем, что файл лежит в папке data, которая находится на одном уровне с src
    # ../data/TK_Uz_Clean.docx
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    FILE_PATH = os.path.join(base_dir, "data", "TK_Uz_Clean.docx")
    DB_DIR = "./chroma_db"

    print("--- 🛠️ ЗАПУСК СТРУКТУРНОЙ ИНДЕКСАЦИИ ---")

    if os.path.exists(DB_DIR):
        print(f"🧹 Удаление старой базы...")
        shutil.rmtree(DB_DIR)

    if not os.path.exists(FILE_PATH):
        print(f"❌ ОШИБКА: Файл не найден по пути: {FILE_PATH}")
        print("Проверь, что в папке 'data' файл называется именно 'TK_Uz_Clean.docx'")
        return

    print(f"📄 Чтение файла из папки data...")
    raw_text = load_docx(FILE_PATH)
    
    print("✂️ Разделение на статьи (Structural Chunking)...")
    docs = split_by_articles(raw_text)
    print(f"✨ Успешно извлечено фрагментов: {len(docs)}")

    print("🧠 Генерация векторов...")
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )

    print("💾 Сохранение в Chroma...")
    db = Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory=DB_DIR
    )
    
    print(f"✅ ГОТОВО! Статей в базе: {len(docs)}")

if __name__ == "__main__":
    rebuild_database()